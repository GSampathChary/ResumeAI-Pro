from __future__ import annotations

from pathlib import Path

from app.parsers.text_cleaner import clean_text


def extract_docx_text(path: str | Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - fallback for missing dependency
        raise RuntimeError("DOCX parsing requires python-docx to be installed") from exc

    document = Document(str(path))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)
    return clean_text("\n".join(chunks))

